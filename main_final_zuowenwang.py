import re
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import shutil


def remove_header_footer(doc, save_doc):
    # doc：需要去页眉页脚的docx 文件
    # finish_doc： 需要另存为的新文件名
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    document.save(save_doc)


def change_word_font(doc_file):
    # 打开doc文件
    doc = Document(doc_file)
    doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
    doc.save(doc_file)


def change_line_spacing(doc_file):
    doc = Document(doc_file)
    for p in doc.paragraphs:  # 循环处理每个段落
        p.paragraph_format.line_spacing = 1.5  # 行距设置为3
    doc.save(doc_file)


number_dirs_arr = [
    0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27
]

if __name__ == '__main__':
    root_dir = "../www.zuowenwang.net"
    for number_dir in number_dirs_arr:
        number_dir = str(number_dir)
        files = sorted(os.listdir(root_dir + "/" + number_dir))
        for file in files:
            if os.path.splitext(file)[1] == ".docx":
                file_path = root_dir + "/" + number_dir + "/" + file
                print(file_path)
                file_finish_dir = root_dir + "/" + number_dir + "/"
                file_finish_dir = file_finish_dir.replace("www.zuowenwang.net", "finish.zuowenwang.net")
                if not os.path.exists(file_finish_dir):
                    os.makedirs(file_finish_dir)
                file_finish_path = file_finish_dir + file.split("-")[1]

                try:
                    # 删除页眉页脚
                    remove_header_footer(file_path, file_finish_path)

                    # 改变文档字体
                    change_word_font(file_finish_path)

                    # 修改行距
                    change_line_spacing(file_finish_path)
                except Exception as e:
                    print(e)
