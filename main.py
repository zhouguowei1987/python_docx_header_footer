import re
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def remove_header_footer(doc, finish_doc):
    # doc：需要去页眉页脚的docx 文件
    # finish_doc： 需要另存为的新文件名
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    document.save(finish_doc)


def get_word_pages(in_file):
    pages = 1
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            word.ActiveDocument.Repaginate()
            pages = word.ActiveDocument.ComputeStatistics(2)
            doc.Close()
            word.Quit()
            return pages
        except Exception as e:
            print(e)
        finally:
            return pages
    except Exception as e:
        print(e)
    finally:
        return pages


def docx_remove_content(doc_file):
    # 定义需要去除的内容
    content_to_remove = '''不用注册，免费下载！'''
    # 打开doc文件
    doc = Document(doc_file)
    # 遍历doc文件中的段落
    for para in doc.paragraphs:
        # 如果段落中包含需要去除的内容，使用正则表达式替换为空字符串
        if re.search(content_to_remove, para.text):
            para.text = re.sub(content_to_remove, '', para.text)

    doc.save(doc_file)


def change_word_font(doc_file):
    # 打开doc文件
    doc = Document(doc_file)
    doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
    doc.save(doc_file)


def check_only_image(doc_file):
    doc = Document(doc_file)
    i = 0
    for para in doc.paragraphs:
        if i == 0 and para.text == "":
            return True
        i += 1
    return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)


if __name__ == '__main__':
    subject_dirs_arr = ['语文试卷', '数学试卷', '英语试卷', '物理试卷', '化学试卷', '政治试卷', '历史试卷', '地理试卷',
                        '生物试卷']
    root_dir = "G:\\final-www.shijuan1.com"
    subject_dirs = sorted(os.listdir(root_dir))
    for subject in subject_dirs:
        if subject in subject_dirs_arr:
            shijuan_dirs = sorted(os.listdir(root_dir + "\\" + subject))

            for shijuan in shijuan_dirs:

                # 删除扫描版
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     if file.find("扫描") != -1:
                #         os.remove(root_dir + "\\" + subject + "\\" + shijuan + "\\" + file)

                # 删除只包含图片
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     print(docx_file)
                #     if check_only_image(docx_file):
                #         os.remove(docx_file)

                # 删除“不用注册，免费下载！”文字
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     print(docx_file)
                #     docx_remove_content(docx_file)

                # 删除“不用注册，免费下载！”文字
                finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                files = sorted(os.listdir(finish_dir))
                for file in files:
                    docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                    print(docx_file)
                    change_word_font(docx_file)

                # 删除文件标题“（word版）”文字
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("（word版）") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("（word版）", "")
                #         os.rename(src_docx_file, dst_docx_file)

                # 替换文件标题“（word解析版）”文字为“（解析版）”
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("（word解析版）") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("（word解析版）", "（解析版）")
                #         os.rename(src_docx_file, dst_docx_file)

                # 替换文件标题“及答案”文字为“（含答案）”
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("及答案") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("及答案", "（含答案）")
                #         os.rename(src_docx_file, dst_docx_file)

                # 删除文件标题“（原卷版）”文字
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("（原卷版）") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("（原卷版）", "")
                #         if not os.path.exists(dst_docx_file):
                #             os.rename(src_docx_file, dst_docx_file)
                #         else:
                #             os.remove(src_docx_file)

                # 替换文件标题“（含解析版）”文字为“（含解析）”
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("（含解析版）") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("（含解析版）", "（含解析）")
                #         os.rename(src_docx_file, dst_docx_file)

                # 将文件标题“（含答案）”文字移到文件名最后位置
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("（含答案）") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("（含答案）", "")
                #         dst_docx_file = dst_docx_file.replace(".docx", "") + "（含答案）.docx"
                #         os.rename(src_docx_file, dst_docx_file)

                # 替换文件标题“解析（含答案）”文字为“（含答案）”
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan
                # files = sorted(os.listdir(finish_dir))
                # for file in files:
                #     src_docx_file = root_dir + "\\" + subject + "\\" + shijuan + "\\" + file
                #     if file.find("解析（含答案）") != -1:
                #         print(src_docx_file)
                #         dst_docx_file = src_docx_file.replace("解析（含答案）", "（含答案）")
                #         if not os.path.exists(dst_docx_file):
                #             os.rename(src_docx_file, dst_docx_file)
                #         else:
                #             os.remove(src_docx_file)

                # if shijuan.find("_finish") != -1 or shijuan.find("_doc2docx") != -1:
                #     continue
                # word_dir = root_dir + "\\" + subject + "\\" + shijuan
                # finish_dir = root_dir + "\\" + subject + "\\" + shijuan + "_finish"
                # doc2docx_dir = root_dir + "\\" + subject + "\\" + shijuan + "_doc2docx"
                #
                # if not os.path.exists(finish_dir):
                #     os.mkdir(finish_dir)
                #
                # if not os.path.exists(doc2docx_dir):
                #     os.mkdir(doc2docx_dir)
                #
                # files = sorted(os.listdir(word_dir))
                # for file in files:
                #     if os.path.splitext(file)[1] in [".doc", ".docx"]:
                #         print(file)
                #         # 查看一下是否已经处理完成
                #         if os.path.isfile(finish_dir + "\\" + os.path.splitext(file)[0] + ".docx"):
                #             continue
                #
                #         if os.path.splitext(file)[1] == ".docx":
                #             # 将文件复制到doc2docx_dir目录
                #             shutil.copyfile(word_dir + "\\" + file, doc2docx_dir + "\\" + file)
                #         elif os.path.splitext(file)[1] == ".doc":
                #             # 将doc文件转化为docx文件
                #             doc2docx(word_dir + "\\" + file, doc2docx_dir + "\\" + os.path.splitext(file)[0] + ".docx")
                #
                #         # 去除word页眉和页脚
                #         doc2docx_file = doc2docx_dir + "\\" + os.path.splitext(file)[0] + ".docx"
                #         finish_doc = finish_dir + "\\" + os.path.splitext(file)[0] + ".docx"
                #         if get_word_pages(doc2docx_file) >= 3:
                #             remove_header_footer(doc2docx_file, finish_doc)
                #             print("=============完成======================")
