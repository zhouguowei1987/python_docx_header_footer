import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    document.save(doc)


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            i = 0
            for para in doc.paragraphs:
                if i == 4 and para.text == "":
                    doc.save(doc_file)
                    return True
                i += 1
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


if __name__ == '__main__':
    category_dirs_arr = ['中考试卷', '高考试卷']
    root_dir = "G:\\www.tc260.org.cn\\www.tc260.org.cn"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)

        docx_dir = "G:\\www.tc260.org.cn\\docx.www.tc260.org.cn" + "\\"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        finish_dir = "G:\\www.tc260.org.cn\\finish.www.tc260.org.cn" + "\\"
        if not os.path.exists(finish_dir):
            os.makedirs(finish_dir)

        fileExt = os.path.splitext(file)[1]
        if fileExt == ".pdf":
            # 是pdf文件，直接复制到finish_dir目录
            shutil.copy(file_path, finish_dir + "\\" + file)
            continue
        else:
            docx_file = docx_dir + "\\" + file.replace(os.path.splitext(file)[1], ".docx")
            if fileExt == ".doc":
                # 将doc文件转化为docx文件
                if not os.path.exists(docx_file):
                    with open(docx_file, 'w') as f:
                        pass
                    print("==========开始转化为docx==============")
                    if not doc2docx(file_path, docx_file):
                        continue
                    print("==========转化完成==============")
            else:
                # 已经是docx文件了，直接复制过去
                shutil.copy(file_path, docx_file)
            # 将docx文件转化为pdf
            finish_file = docx_file.replace("docx.", "finish.").replace(".docx", ".pdf")
            if not os.path.exists(finish_file):
                # 将docx转化为pdf
                with open(finish_file, "w") as f:
                    # 将 Word 文档转换为 PDF
                    try:
                        print("==========开始转化为pdf==============")
                        convert(docx_file, finish_file)
                        print("转换成功！")
                    except Exception as e:
                        print("转换失败：", str(e))
