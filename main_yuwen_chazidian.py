import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil
import rarfile
import re

rarfile.UNRAR_TOOL = "D:\\Program Files (x86)\\WinRAR\\UnRAR.exe"


def change_word_font(doc_file):
    try:
        # 打开doc文件
        doc = Document(doc_file)
        doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    try:
        document = Document(doc)
        for section in document.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        document.save(doc)
        return True
    except Exception as e:
        print(e)
        return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            i = 0
            for para in doc.paragraphs:
                if i == 4 and para.text == "":
                    doc.save(doc_file)
                    return True
                i += 1
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


def is_pure_number(string):
    pattern = r'^[0-9]+$'
    if re.match(pattern, string):
        return True
    else:
        return False


def decompress_rar(rar_file_name, dir_name):
    """
    .rar 文件解压
    :param rar_file_name: rar 文件路径
    :param dir_name: 文件解压目录
    :return:
    """
    # 创建 rar 对象
    rar_obj = rarfile.RarFile(rar_file_name)
    # 目录切换
    if not os.path.exists(dir_name):
        os.mkdir(dir_name)
    os.chdir(dir_name)
    # Extract all files into current directory.
    rar_obj.extractall()
    # rar_obj.extractall(dir_name)
    # 关闭
    rar_obj.close()


if __name__ == '__main__':
    # 解压压缩包
    # rar_root_dir = "E:\\workspace\\yuwen.chazidian.com\\yuwen.rar_chazidian.com"
    # rar_dirs = sorted(os.listdir(rar_root_dir))
    # rar_files = sorted(os.listdir(rar_root_dir))
    # for rar_file in rar_files:
    #     rar_file_path = rar_root_dir + "\\" + rar_file
    #     dst_file_path = "E:\\workspace\\yuwen.chazidian.com\\yuwen.uncompress_chazidian.com"
    #     if os.path.splitext(rar_file)[1] in [".doc", ".docx"]:
    #         print("==========" + "开始复制" + "==========")
    #         shutil.copy(rar_file_path, dst_file_path + "\\" + rar_file)
    #         print("==========" + "复制完成" + "==========")
    #     else:
    #         print("==========" + "开始解压" + rar_file_path + "==========")
    #         try:
    #             decompress_rar(rar_file_path, dst_file_path)
    #         except Exception as e:
    #             print(e)
    #             continue
    #         print("==========" + "解压完成" + "==========")
    # exit()

    root_dir = "E:\\workspace\\yuwen.chazidian.com\\yuwen.uncompress_chazidian.com"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)

        # 查看一下是否是文件夹，如果是文件夹，则将文件移出
        # if os.path.isdir(file_path):
        #     child_files = sorted(os.listdir(file_path))
        #     for child_file in child_files:
        #         extension = os.path.splitext(child_file)[-1]
        #         if extension not in [".doc", ".docx"]:
        #             continue
        #         src_child_file_path = file_path + "\\" + child_file
        #         dst_child_file_path = root_dir + "\\" + child_file
        #         try:
        #             os.rename(src_child_file_path, dst_child_file_path)
        #         except WindowsError:
        #             os.remove(dst_child_file_path)
        #             os.rename(src_child_file_path, dst_child_file_path)

        # 文件名长度小于20，则删除
        # if len(file) < 20:
        #     os.remove(file_path)

        # 文件后缀不是doc或docx，则删除
        # if os.path.splitext(file)[1] not in [".doc", ".docx"]:
        #     os.remove(file_path)

        # 文件名是纯数字删除
        # if is_pure_number(os.path.splitext(file)[0]):
        #     os.remove(file_path)

        docx_dir = "E:\\workspace\\yuwen.chazidian.com\\yuwen.docx_chazidian.com"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        sub_file = file
        left_flag_index = file.find("【")
        right_flag_index = file.find("】")
        if left_flag_index == 0 and right_flag_index != -1:
            # 文档名称以“【”开头，以“】”结尾，则替换名称
            sub_file = file[right_flag_index + 1:]
        docx_file = docx_dir + "\\" + sub_file.lower().replace(os.path.splitext(sub_file)[1], ".docx")
        docx_file = docx_file.strip()
        docx_file = docx_file.replace("（", "(").replace("）", ")")
        docx_file = docx_file.replace("＜＜", "(").replace("＞＞", ")")
        docx_file = docx_file.replace("[", "(").replace("]", ")")
        docx_file = docx_file.replace("(word版)", "")
        docx_file = docx_file.replace("zhoushile-", "")
        docx_file = docx_file.replace(",", "")
        docx_file = docx_file.replace("，", "")

        # 删除文件名中不含有“数学”字样文件
        if "语文" not in docx_file:
            os.remove(file_path)
            continue

        # if not os.path.exists(docx_file):
        #     # 获取文件后缀
        #     file_ext = os.path.splitext(file_path)[-1]
        #     if file_ext == ".docx":
        #         # 已经是docx文件了，直接复制过去
        #         shutil.copy(file_path, docx_file)
        #     else:
        #         with open(docx_file, 'w') as f:
        #             pass
        #         print("==========开始转化为docx==============")
        #         if not doc2docx(file_path, docx_file):
        #             # 删除原文件
        #             os.remove(file_path)
        #             os.remove(docx_file)
        #             continue
        #         print("==========转化完成==============")
        #
        # if os.path.exists(docx_file):
        #     # 删除只包含图片
        #     if check_only_image(docx_file):
        #         # 删除原文件
        #         os.remove(file_path)
        #         # 删除图片文件
        #         os.remove(docx_file)
        #         continue
        #
        #     # 删除页眉页脚
        #     if not remove_header_footer(docx_file):
        #         # 删除原文件
        #         os.remove(file_path)
        #         os.remove(docx_file)
        #         continue
        #
        #     # 改变文档字体
        #     if not change_word_font(docx_file):
        #         # 删除原文件
        #         os.remove(file_path)
        #         os.remove(docx_file)
        #         continue

        # docx文件已存在，跳过继续
        if os.path.exists(docx_file):
            # continue
            finish_dir = "E:\\workspace\\yuwen.chazidian.com\\yuwen.finish_chazidian.com"
            if not os.path.exists(finish_dir):
                os.makedirs(finish_dir)
            # 将docx文件转化为pdf
            finish_file = docx_file.replace("docx_", "finish_").replace(".docx", ".pdf")
            if not os.path.exists(finish_file):
                # 将docx转化为pdf
                with open(finish_file, "w") as f:
                    # 将 Word 文档转换为 PDF
                    try:
                        print("==========开始转化为pdf==============")
                        convert(docx_file, finish_file)
                        print("转换成功！")
                    except Exception as e:
                        print("转换失败：", str(e))

