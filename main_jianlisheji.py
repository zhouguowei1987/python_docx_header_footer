import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def change_word_font(doc_file):
    try:
        # 打开doc文件
        doc = Document(doc_file)
        doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def docx_remove_content(doc_file):
    try:
        # 定义需要去除及替换的内容
        content_to_removes = [
            ['''简历设计网''', '个人简历'],
            ['''jianlisheji''', 'XXXXXX'],
            ['''jinalisheji''', 'XXXXXX'],
            ['''Jianlisheji''', 'XXXXXX'],
            ['''Jianliseji''', 'XXXXXX'],
        ]
        # 打开doc文件
        doc = Document(doc_file)

        # 遍历文本框
        for i in range(len(doc.inline_shapes._body.xpath('//w:txbxContent'))):
            for tx_idx, tx in enumerate(doc.inline_shapes._body.xpath('//w:txbxContent')[i]):
                children = tx.getchildren()
                for child_idx, child in enumerate(children):
                    if child.text:
                        for content_to_remove in content_to_removes:
                            target_text = content_to_remove[0]
                            replacement_text = content_to_remove[1]
                            if target_text in child.text:
                                child.text = child.text.replace(target_text, replacement_text)

        # 遍历表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 遍历表格段落内容，回到上个步骤，将cell当作paragraph处理
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # 替换功能
                            for content_to_remove in content_to_removes:
                                if content_to_remove[0] in cell.text:
                                    run.text = run.text.replace(content_to_remove[0], content_to_remove[1])

        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    try:
        document = Document(doc)
        for section in document.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        document.save(doc)
        return True
    except Exception as e:
        print(e)
        return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


if __name__ == '__main__':
    root_dir = "F:\\workspace\\www.jianlisheji.com\\www.jianlisheji.com"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)
        docx_dir = "F:\\workspace\\www.jianlisheji.com\\docx.jianlisheji.com"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        docx_file = docx_dir + "\\" + file.replace(os.path.splitext(file)[1], ".docx")
        print(docx_file)

        # 获取文件后缀
        file_ext = os.path.splitext(file_path)[-1]
        if file_ext == ".docx":
            # 已经是docx文件了，直接复制过去
            shutil.copy(file_path, docx_file)
        else:
            with open(docx_file, 'w') as f:
                pass
            print("==========开始转化为docx==============")
            if not doc2docx(file_path, docx_file):
                os.remove(docx_file)
                continue
            print("==========转化完成==============")

        if os.path.exists(docx_file):
            # 改变文档文字
            if not docx_remove_content(docx_file):
                # 删除原文件
                # os.remove(file_path)
                os.remove(docx_file)
                continue

            # 删除页眉页脚
            if not remove_header_footer(docx_file):
                # 删除原文件
                # os.remove(file_path)
                os.remove(docx_file)
                continue

            # 改变文档字体
            if not change_word_font(docx_file):
                # 删除原文件
                # os.remove(file_path)
                os.remove(docx_file)
                continue
