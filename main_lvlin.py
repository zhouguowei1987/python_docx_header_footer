import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
import os
import shutil


def check_header(doc):
    # doc：需要去页眉页脚的docx 文件
    document = Document(doc)
    if not document.sections[0].header.is_linked_to_previous:
        return True
    return False


def remove_header_footer(doc_file):
    # doc：需要去页眉页脚的docx 文件
    doc = Document(doc_file)
    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    doc.save(doc_file)


def start_copy_paragraph(doc_file):
    try:
        doc = Document(doc_file)
        paragraphs = doc.paragraphs
        # 从哪个段落开始复制
        start_copy_paragraph_index = len(paragraphs) - 1
        # 最后一个段落
        para = paragraphs[len(paragraphs) - 1]
        p = para._element
        img = p.xpath('.//pic:pic')
        if img:
            # 最后一个段落是图片，跳过（删除）
            # p.getparent().remove(p)
            # para._p = para._element = None
            start_copy_paragraph_index = len(paragraphs) - 2

        new_doc = Document()
        new_doc.styles['Normal'].font.name = 'Times New Roman'
        new_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        for para in doc.paragraphs[:start_copy_paragraph_index]:
            new_para = new_doc.add_paragraph()
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.paragraph_format.keep_together = para.paragraph_format.keep_together
            new_para.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            new_para.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.space_after = para.paragraph_format.space_after
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            for run in para.runs:
                output_run = new_para.add_run(run.text)
                output_run.bold = run.bold
                output_run.italic = run.italic
                output_run.underline = run.underline
                output_run.font.size = run.font.size
                output_run.font.color.rgb = run.font.color.rgb
                output_run.style.name = run.style.name
        new_doc.save(doc_file)
    except Exception as e:
        print(e)
        print("删除文件")
        os.remove(doc_file)
        return True
    return False


if __name__ == '__main__':
    root_dir = "../lvlin.baidu.com/lvlin.baidu.com"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "/" + file
        print(file_path)
        docx_dir = "../lvlin.baidu.com/docx.lvlin.baidu.com"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        docx_file = docx_dir + "/" + file

        # 获取文件后缀
        file_ext = os.path.splitext(file_path)[-1]
        if file_ext == ".docx":
            # 已经是docx文件了，直接复制过去
            shutil.copy(file_path, docx_file)

        if os.path.exists(docx_file):
            # 开始复制文档
            start_copy_paragraph(docx_file)
