import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def change_word_font(doc_file):
    try:
        # 打开doc文件
        doc = Document(doc_file)
        doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def convertDocxToPDF(infile, outfile):
    wdFormatPDF = 17
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            # 连续5个段落都是空，则按照纯图片处理
            if doc.paragraphs[0].text == "" and doc.paragraphs[1].text == "" and doc.paragraphs[2].text == "" and \
                    doc.paragraphs[3].text == "" and doc.paragraphs[4].text == "":
                doc.save(doc_file)
                return True
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


def remove_header_footer(doc_file):
    # doc：需要去页眉页脚的docx 文件
    try:
        doc = Document(doc_file)
        for section in doc.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            if doc.ProtectionType == 1:
                print('文档加密，转换失败')
                doc.Close()
                word.Quit()
                return False
            else:
                doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
                print('转换成功')
                doc.Close()
                word.Quit()
                return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


if __name__ == '__main__':
    category_dirs_arr = ['语文', '英语', '数学', '科学', '道德与法治']
    # category_dirs_arr = ['语文', '数学', '英语','道德与法治','地理','化学','科学','历史','历史与社会','美术','生物','体育与健康','物理','信息技术','音乐','综合']
    # category_dirs_arr = ['语文', '数学','英语', '物理', '化学', '生物', '通用技术', '信息技术', '综合', '政治', '地理', '历史']
    root_dir = "E:\\workspace\\www.zhuangyuan123.com\\www.zhuangyuan123.com\\小学"
    category_dirs = sorted(os.listdir(root_dir))
    for category in category_dirs:
        if category in category_dirs_arr:
            files = sorted(os.listdir(root_dir + "\\" + category))
            for file in files:
                print(file)

                file_path = root_dir + "/" + category + "/" + file
                print(file_path)

                # 直接使用splitext并忽略扩展名部分
                filename_without_ext = os.path.splitext(file_path)[0]
                # 如果文件结尾是数字，则删除
                last_letter = filename_without_ext[-1]
                if last_letter in [".", "0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]:
                    print("删除文件")
                    os.remove(file_path)
