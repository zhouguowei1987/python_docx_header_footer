import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
from pptx import Presentation
import os
import shutil


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    try:
        document = Document(doc)
        for section in document.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        document.save(doc)
        return True
    except Exception as e:
        print(e)
        return False


def docx_remove_content(doc_file):
    # 定义需要去除及替换的内容
    content_to_removes = [
        ['''法大大''', '小编']
    ]
    # 打开doc文件
    doc = Document(doc_file)
    doc.paragraphs[1].clear()
    # 遍历doc文件中的段落
    for para in doc.paragraphs:
        # 如果段落中包含需要去除的内容，使用正则表达式
        for content_to_remove in content_to_removes:
            if re.search(content_to_remove[0], para.text):
                para.text = re.sub(content_to_remove[0], content_to_remove[1], para.text)

    doc.save(doc_file)


def change_word_font(doc_file):
    # 打开doc文件
    doc = Document(doc_file)
    doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体

    # 修改标题字体
    para = doc.paragraphs[0]
    for run in para.runs:
        # run.font.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(doc_file)


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 16, False, "", True, "", False, False, False, False)
            # doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
            return False
    except Exception as e:
        print(e)
    return False


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            i = 0
            for para in doc.paragraphs:
                if i == 4 and para.text == "":
                    doc.save(doc_file)
                    return True
                i += 1
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


if __name__ == '__main__':
    # 将doc文档转化为docx
    root_dir = "E:\\workspace\\m.fadada.com\\temp-m.fadada.com"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)
        docx_dir = "E:\\workspace\\m.fadada.com\\docx.m.fadada.com"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        sub_file = file
        left_flag_index = file.find("【")
        right_flag_index = file.find("】")
        if left_flag_index == 0 and right_flag_index != -1:
            # 文档名称以“【”开头，以“】”结尾，则替换名称
            sub_file = file[right_flag_index + 1:]
        sub_file = os.path.splitext(sub_file)[0].replace(".doc", "").replace(".docx", "") + ".docx"
        docx_file = docx_dir + "\\" + sub_file
        docx_file = docx_file.replace(" ", "")
        docx_file = docx_file.replace("（", "(").replace("）", ")")
        docx_file = docx_file.replace("[", "").replace("]", "")
        docx_file = docx_file.replace("()", "")
        docx_file = docx_file.replace("-", "")
        docx_file = docx_file.replace("——", "-")
        docx_file = docx_file.replace(",", "")
        docx_file = docx_file.replace("，", "")
        print(docx_file)

        if not os.path.exists(docx_file):
            with open(docx_file, 'w') as f:
                pass
            print("==========开始转化为docx==============")
            if not doc2docx(file_path, docx_file):
                # 删除原文件
                os.remove(file_path)
                os.remove(docx_file)
                continue
            print("==========转化完成==============")
            # 获取文件后缀
            file_ext = os.path.splitext(file_path)[-1]
            if file_ext == ".docx":
                try:
                    # 已经是docx文件了，直接复制过去
                    shutil.copy(file_path, docx_file)
                    print("File copied successfully.")
                except FileNotFoundError:
                    print("The source file does not exist.")
                except PermissionError:
                    print("Permission denied.")
                except shutil.SameFileError:
                    print("The source and destination are the same file.")
                except shutil.Error as e:
                    print(f"An error occurred: {e}")
            else:
                with open(docx_file, 'w') as f:
                    pass
                print("==========开始转化为docx==============")
                if not doc2docx(file_path, docx_file):
                    # 删除原文件
                    os.remove(file_path)
                    os.remove(docx_file)
                    continue
                print("==========转化完成==============")

        if os.path.exists(docx_file):
            # 删除只包含图片
            if check_only_image(docx_file):
                # 删除原文件
                os.remove(file_path)
                os.remove(docx_file)
                continue

            # 删除页眉页脚
            if not remove_header_footer(docx_file):
                # 删除原文件
                os.remove(file_path)
                os.remove(docx_file)
                continue
            # 过滤文档文字
            docx_remove_content(docx_file)

            # 改变文档字体
            change_word_font(docx_file)
        # 删除temp文件夹文件
        os.remove(file_path)
