import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def remove_and_set_header_footer(doc, save_doc):
    # doc：需要去页眉页脚的docx 文件
    # finish_doc： 需要另存为的新文件名
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    document.sections[0].header_distance = Cm(1.5)
    document.sections[0].footer_distance = Cm(1.75)
    # 设置页眉
    header = document.sections[0].header  # 获取第一个节的页眉（所有的页眉都一致）
    paragraph = header.paragraphs[0]  # 获取页眉的文字part
    text = paragraph.add_run('烦恼多多少少，放松必不可少；给自己一个微笑，迎来的将是一片美好！')
    text.font.size = Pt(10)  # 页眉字体大小
    text.font.color.rgb = RGBColor(255, 0, 0)
    # text.bold = True  # 页眉字体是否加粗
    text.font.name = 'Times New Roman'  # 控制是英文时的字体
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 控制是中文时的字体
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置页眉居中

    document.save(save_doc)


def get_word_pages(in_file):
    pages = 0
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            word.ActiveDocument.Repaginate()
            pages = word.ActiveDocument.ComputeStatistics(2)
            doc.Close()
            word.Quit()
            return pages
        except Exception as e:
            print(e)
        finally:
            return pages
    except Exception as e:
        print(e)
    finally:
        return pages


def docx_get_word_pages(in_file):
    pages = 0
    try:
        # 打开Word文档
        doc = Document(in_file)
        # 获取总页数
        pages = len(doc.paragraphs)
        return pages
    except Exception as e:
        print(e)
    finally:
        return pages


def docx_remove_content(doc_file):
    # 定义需要去除的内容
    content_to_remove = '''XXXXXX'''
    # 打开doc文件
    doc = Document(doc_file)
    # 遍历doc文件中的段落
    for para in doc.paragraphs:
        # 如果段落中包含需要去除的内容，使用正则表达式替换为空字符串
        if re.search(content_to_remove, para.text):
            para.text = re.sub(content_to_remove, 'OfficePLUS', para.text)

    doc.save(doc_file)


def change_word_font(doc_file):
    # 打开doc文件
    doc = Document(doc_file)
    doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体

    i = 0
    for para in doc.paragraphs:
        if i == 0:
            for run in para.runs:
                # run.font.bold = True
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(255, 0, 0)
        i += 1
    doc.save(doc_file)


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            i = 0
            for para in doc.paragraphs:
                if i == 4 and para.text == "":
                    doc.save(doc_file)
                    return True
                i += 1
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


if __name__ == '__main__':
    category_dirs_arr = ['自考', '专升本考试', '小升初', '考研', '高中会考', '高考', '成人高考', '中考']
    root_dir = "G:\\tikuvip（2016-2017）.51test.net"
    category_dirs = sorted(os.listdir(root_dir))
    for category in category_dirs:
        if category in category_dirs_arr:
            files = sorted(os.listdir(root_dir + "\\" + category))
            for file in files:
                if os.path.splitext(file)[1] == ".doc":
                    file_path = root_dir + "\\" + category + "\\" + file
                    print(file_path)
                    if "答案" not in file:
                        continue

                    docx_dir = "G:\\docx.tikuvip（2016-2017）.51test.net" + "\\" + category
                    if not os.path.exists(docx_dir):
                        os.makedirs(docx_dir)

                    docx_file = docx_dir + "\\" + file.replace(".doc", ".docx")
                    if not os.path.exists(docx_file):
                        print("==========开始转化为docx==============")
                        if not doc2docx(file_path, docx_file):
                            continue
                        print("==========转化完成==============")

                    finish_dir = "G:\\finish.tikuvip（2016-2017）.51test.net" + "\\" + category
                    if not os.path.exists(finish_dir):
                        os.makedirs(finish_dir)
                    finish_file = finish_dir + "\\" + file.replace(".doc", ".pdf")

                    # replace_text = "(含答案)"
                    # if "及答案" in file:
                    #     finish_file = os.path.splitext(finish_file)[0].replace("及答案", "") + replace_text + \
                    #                   os.path.splitext(finish_file)[1]
                    # if "与答案" in file:
                    #     finish_file = os.path.splitext(finish_file)[0].replace("与答案", "") + replace_text + \
                    #                   os.path.splitext(finish_file)[1]
                    # if "含答案" in file:
                    #     finish_file = os.path.splitext(finish_file)[0].replace("含答案", "") + replace_text + \
                    #                   os.path.splitext(finish_file)[1]
                    # if "附答案" in file:
                    #     finish_file = os.path.splitext(finish_file)[0].replace("附答案", "") + replace_text + \
                    #                   os.path.splitext(finish_file)[1]
                    if not os.path.exists(finish_file):
                        # 删除只包含图片
                        if check_only_image(docx_file):
                            continue
                        # 将docx转化为pdf
                        with open(finish_file, "w") as f:
                            # 将 Word 文档转换为 PDF
                            try:
                                convert(docx_file, finish_file)
                                print("转换成功！")
                            except Exception as e:
                                print("转换失败：", str(e))

                        # shutil.copy(docx_file, finish_file)
                        # 删除并设置页眉页脚
                        # remove_and_set_header_footer(docx_file, finish_file)
                        # 改变文档字体
                        # change_word_font(finish_file)
