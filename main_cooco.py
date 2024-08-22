import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def change_word_font(doc_file):
    try:
        # 打开doc文件
        doc = Document(doc_file)
        doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            # 连续5个段落都是空，则按照纯图片处理
            if doc.paragraphs[0].text == "" and doc.paragraphs[1].text == "" and doc.paragraphs[2].text == "" and \
                    doc.paragraphs[3].text == "" and doc.paragraphs[4].text == "":
                doc.save(doc_file)
                return True
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


def remove_header_footer(doc_file):
    # doc：需要去页眉页脚的docx 文件
    try:
        doc = Document(doc_file)
        for section in doc.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            if doc.ProtectionType == 1:
                print('文档加密，转换失败')
                doc.Close()
                word.Quit()
                return False
            else:
                doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
                print('转换成功')
                doc.Close()
                word.Quit()
                return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


if __name__ == '__main__':
    # category_dirs_arr = ['语文', '数学', '英语', '物理', '化学', '地理', '历史', '生物', '政治']
    # category_dirs_arr = ['语文', '数学', '英语', '物理', '化学', '地理', '历史', '生物', '政治']
    category_dirs_arr = ['语文', '数学', '英语', '政治']
    root_dir = "E:\\workspace\\bk.cooco.net.cn\\bk.cooco.net.cn\\小学"
    category_dirs = sorted(os.listdir(root_dir))
    for category in category_dirs:
        if category in category_dirs_arr:
            files = sorted(os.listdir(root_dir + "\\" + category))
            for file in files:
                print(file)

                file_path = root_dir + "/" + category + "/" + file
                print(file_path)
                docx_dir = "E:\\workspace\\bk.cooco.net.cn\\bk.docx_cooco.net.cn\\小学\\" + category
                if not os.path.exists(docx_dir):
                    os.makedirs(docx_dir)

                sub_file = file
                left_flag_index = file.find("【")
                right_flag_index = file.find("】")
                if left_flag_index == 0 and right_flag_index != -1:
                    # 文档名称以“【”开头，以“】”结尾，则替换名称
                    sub_file = file[right_flag_index + 1:]
                docx_file = docx_dir + "\\" + sub_file.lower().replace(os.path.splitext(sub_file)[1], ".docx")
                docx_file = docx_file.replace(" ", "")
                docx_file = docx_file.replace("word版，", "")
                docx_file = docx_file.replace("word版", "")
                docx_file = docx_file.replace("有答案", "含答案")
                docx_file = docx_file.replace("(无答案)", "")
                docx_file = docx_file.replace("(免费)", "")
                docx_file = docx_file.replace("()", "")
                docx_file = docx_file.replace("——", "-")

                # 标题含有扫描字样
                if docx_file.find("扫描") != -1:
                    # 文档标题不包含分类名称
                    print("标题含有扫描字样，跳过")
                    continue

                if docx_file.replace(docx_dir, "").find(category) == -1:
                    # 文档标题不包含分类名称
                    print("文档标题不包含分类名称，跳过")
                    continue

                # docx文件已存在，跳过继续
                if not os.path.exists(docx_file):
                    # 获取文件后缀
                    file_ext = os.path.splitext(file_path)[-1]
                    if file_ext == ".docx":
                        # 已经是docx文件了，直接复制过去
                        shutil.copy(file_path, docx_file)
                    else:
                        with open(docx_file, 'w') as f:
                            pass
                        print("==========开始转化为docx==============")
                        if not doc2docx(file_path, docx_file):
                            print("转换失败，删除文件")
                            os.remove(docx_file)
                            continue
                        print("==========转化完成==============")

                if os.path.exists(docx_file):

                    # 删除页眉页脚
                    if not remove_header_footer(docx_file):
                        print("删除页眉页脚失败，删除文件")
                        os.remove(docx_file)
                        continue

                    # 删除只包含图片
                    if check_only_image(docx_file):
                        print("文档只包含图片，删除文件")
                        os.remove(docx_file)
                        continue

                    # 改变文档字体
                    if not change_word_font(docx_file):
                        print("改变文档字体失败，删除文件")
                        os.remove(docx_file)
                        continue

                # docx文件已存在，跳过继续
                if os.path.exists(docx_file):
                    # continue
                    finish_dir = "E:\\workspace\\bk.cooco.net.cn\\bk.finish_cooco.net.cn\\小学\\" + category
                    if not os.path.exists(finish_dir):
                        os.makedirs(finish_dir)
                    # 将docx文件转化为pdf
                    finish_file = docx_file.replace("docx_", "finish_").replace(".docx", ".pdf")
                    if not os.path.exists(finish_file):
                        # 将docx转化为pdf
                        with open(finish_file, "w") as f:
                            # 将 Word 文档转换为 PDF
                            try:
                                print("==========开始转化为pdf==============")
                                convert(docx_file, finish_file)
                                print("转换成功！")
                            except Exception as e:
                                print("转换失败：", str(e))
