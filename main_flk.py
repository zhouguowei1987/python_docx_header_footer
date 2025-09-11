import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
from pptx import Presentation
import os
import shutil


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    try:
        document = Document(doc)
        for section in document.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        document.save(doc)
        return True
    except Exception as e:
        print(e)
        return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 16, False, "", True, "", False, False, False, False)
            # doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
            return False
    except Exception as e:
        print(e)
    return False


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            i = 0
            for para in doc.paragraphs:
                if i == 4 and para.text == "":
                    doc.save(doc_file)
                    return True
                i += 1
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


if __name__ == '__main__':
    # 将doc文档转化为docx
    root_dir = "E:\\workspace\\flk.npc.gov.cn\\temp-flk.npc.gov.cn"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)
        docx_dir = "E:\\workspace\\upload.doc88.com\\flk.npc.gov.cn"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        sub_file = file
        sub_file = os.path.splitext(sub_file)[0].replace(".doc", "").replace(".docx", "") + ".docx"
        docx_file = docx_dir + "\\" + sub_file
        print(docx_file)

        if not os.path.exists(docx_file):
            # 获取文件后缀
            file_ext = os.path.splitext(file_path)[-1]
            if file_ext == ".docx":
                try:
                    # 已经是docx文件了，直接复制过去
                    shutil.copy(file_path, docx_file)
                    print("File copied successfully.")
                except FileNotFoundError:
                    print("The source file does not exist.")
                except PermissionError:
                    print("Permission denied.")
                except shutil.SameFileError:
                    print("The source and destination are the same file.")
                except shutil.Error as e:
                    print(f"An error occurred: {e}")
            else:
                with open(docx_file, 'w') as f:
                    pass
                print("==========开始转化为docx==============")
                if not doc2docx(file_path, docx_file):
                    # 删除原文件
                    os.remove(file_path)
                    os.remove(docx_file)
                    continue
                print("==========转化完成==============")

        if os.path.exists(docx_file):
            # 删除只包含图片
            if check_only_image(docx_file):
                # 删除图片文件
                os.remove(docx_file)
                continue

            # 删除页眉页脚
            if not remove_header_footer(docx_file):
                # 删除原文件
                os.remove(docx_file)
                continue

        # docx文件已存在，跳过继续
        if os.path.exists(docx_file):
            # continue
            finish_dir = "E:\\workspace\\upload.doc88.com\\flk.npc.gov.cn"
            if not os.path.exists(finish_dir):
                os.makedirs(finish_dir)
            # 将docx文件转化为pdf
            finish_file = docx_file.replace(".docx", ".pdf")
            if not os.path.exists(finish_file):
                # 将docx转化为pdf
                with open(finish_file, "w") as f:
                    # 将 Word 文档转换为 PDF
                    try:
                        print("==========开始转化为pdf==============")
                        convert(docx_file, finish_file)
                        print("转换成功！")
                    except Exception as e:
                        print("转换失败：", str(e))
        # 删除temp文件夹文件
        os.remove(file_path)
