import re
import time
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
import os
from win32com import client as wc
from win32com.client import constants  # 导入枚举常数模块
import shutil


def remove_header_footer(doc, save_doc):
    # doc：需要去页眉页脚的docx 文件
    # finish_doc： 需要另存为的新文件名
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    document.save(save_doc)


def docx_remove_content(doc_file):
    # 定义需要去除及替换的内容
    content_to_removes = [
        ['''论文类别：(.*?)
上传时间：(.*?)
论文作者：(.*?)
论文版本：(.*?) (.*?) (.*?)''', ''],
        ['''免费论文下载中 http://www.hi138.com 　　''', '\n\r\t'],
        ['''免费论文下载中心 http://www.hi138.com''', ''],
        ['''免费论文下载中 http://www.hi138.com''', ''],
        ['''免费论文下载中心讯：''', ''],
        ['''代写论文网： ''', ''],
        ['''声明：
本论文来自免费论文下载中心：(.*?)
免费论文下载中心（www.hi138.com）(.*?)，本站仅供大家学习、研究、参考之用，未取得作者授权严禁摘编、篡改、用作商业用途.''', '']
    ]
    # 打开doc文件
    doc = Document(doc_file)
    # doc.paragraphs[1].clear()
    # 遍历doc文件中的段落
    for para in doc.paragraphs:
        # 如果段落中包含需要去除的内容，使用正则表达式
        for content_to_remove in content_to_removes:
            if re.search(content_to_remove[0], para.text):
                para.text = re.sub(content_to_remove[0], content_to_remove[1], para.text)

    doc.save(doc_file)


def change_word_font(doc_file):
    # 打开doc文件
    doc = Document(doc_file)
    doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体

    # 修改标题字体
    para = doc.paragraphs[0]
    for run in para.runs:
        # run.font.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(doc_file)


def delete_blank_line(doc_file):
    doc = Document(doc_file)
    for p in doc.paragraphs:  # 循环处理每个段落
        if len(p.text.replace("\n", "").replace(" ", "")) == 0:
            p.clear()
    doc.save(doc_file)


def change_line_spacing(doc_file):
    doc = Document(doc_file)
    for p in doc.paragraphs:  # 循环处理每个段落
        p.paragraph_format.line_spacing = 1.5  # 行距设置为3
    doc.save(doc_file)


def doc2docx(in_file, out_file):
    returnBool = False
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            # 删除文档中超链接（保留文字）
            hylCount = doc.Hyperlinks.Count  # 文档中超链接的总数
            for j in range(0, hylCount):  # 遍历超链接
                # 是否需要删除文本根据需要选择下列两句中的一句
                # 因为倒着删除比较保险，所以用Hyperlinks(hylCount-j)
                doc.Hyperlinks(hylCount - j).Delete()  # 删除超链接（保留纯文本）
                # doc.Hyperlinks(hylCount-j).Range.Delete() # 删除超链接区域（包括文本全部删除）
            # doc.Close(constants.wdSaveChanges)  # 保存并关闭文件
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            doc.Close()
            returnBool = True
        except Exception as e:
            print(e)
            returnBool = False
        finally:
            word.Quit()
    except Exception as e:
        print(e)
        returnBool = False
    finally:
        return returnBool


if __name__ == '__main__':
    root_dir = "G:\\www.hi138.com\\"
    files = sorted(os.listdir(root_dir))
    for file in files:
        if os.path.splitext(file)[1] == ".doc":
            file_path = root_dir + file
            print(file_path + "==============")

            docx_dir = "G:\\docx.hi138.com\\"
            if not os.path.exists(docx_dir):
                os.mkdir(docx_dir)

            docx_file = docx_dir + file.replace(".doc", ".docx")
            if not os.path.exists(docx_file):
                print("==========开始转化==============")
                if not doc2docx(file_path, docx_file):
                    os.remove(file_path)
                    continue
                print("==========转化完成==============")

            finish_dir = "G:\\finish.hi138.com\\"
            if not os.path.exists(finish_dir):
                os.mkdir(finish_dir)

            finish_file = finish_dir + file.replace(".doc", ".docx")
            if not os.path.exists(finish_file):
                try:
                    print("==========文档处理==============")
                    # 删除页眉页脚
                    remove_header_footer(docx_file, finish_file)

                    # 过滤文档文字
                    docx_remove_content(finish_file)

                    # 删除空白行
                    delete_blank_line(finish_file)

                    # 改变文档字体
                    # change_word_font(finish_file)

                    # 修改行距
                    change_line_spacing(finish_file)
                    print("==========处理完成==============")
                except Exception as e:
                    print(e)
