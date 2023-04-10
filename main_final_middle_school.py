import re
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from win32com import client as wc
import os
import shutil


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    # finish_doc： 需要另存为的新文件名
    document = Document(doc)
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    document.save(doc)


def get_word_pages(in_file):
    pages = 1
    try:
        word = wc.Dispatch("Word.Application")
        try:
            doc = word.Documents.Open(in_file)
            word.ActiveDocument.Repaginate()
            pages = word.ActiveDocument.ComputeStatistics(2)
            doc.Close()
            word.Quit()
            return pages
        except Exception as e:
            print(e)
        finally:
            return pages
    except Exception as e:
        print(e)
    finally:
        return pages


def docx_remove_content(doc_file):
    # 定义需要去除的内容
    content_to_remove = '''XXXXXX'''
    # 打开doc文件
    doc = Document(doc_file)
    # 遍历doc文件中的段落
    for para in doc.paragraphs:
        # 如果段落中包含需要去除的内容，使用正则表达式替换为空字符串
        if re.search(content_to_remove, para.text):
            para.text = re.sub(content_to_remove, 'OfficePLUS', para.text)

    doc.save(doc_file)


def change_word_font(doc_file):
    # 打开doc文件
    doc = Document(doc_file)
    doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
    doc.save(doc_file)


def print_filter_word(doc_file):
    doc = Document(doc_file)
    filter_words = ["来源"]
    stop = False
    dst_doc_file = doc_file
    for word in filter_words:
        for para in doc.paragraphs:
            if word in para.text:
                print("=======" + doc_file + "========")
                filter_docx_file = doc_file.replace("Word模板", "filter-Word模板")
                if not os.path.exists(filter_docx_file):
                    os.rename(doc_file, filter_docx_file)
                dst_doc_file = filter_docx_file
                stop = True
                break
        if stop:
            break
    doc.save(dst_doc_file)


def rename_docx_name(doc_file):
    filter_words = ["参考答案", "答案要点", "试卷答案", "【答案】", "解析"]
    dst_docx_file = doc_file
    stop = False
    doc = Document(doc_file)
    for word in filter_words:
        for para in doc.paragraphs:
            if word in para.text:
                filter_docx_file = doc_file
                if word == "解析":
                    filter_docx_file = doc_file.replace("（含答案）", "") \
                                           .replace("（含解析）", "") \
                                           .replace(".docx", "") + "（含解析）.docx"
                elif word == "参考答案":
                    filter_docx_file = doc_file.replace("（含答案）", "") \
                                           .replace("（含解析）", "") \
                                           .replace(".docx", "") + "（含答案）.docx"
                elif word == "答案要点":
                    filter_docx_file = doc_file.replace("（含答案）", "") \
                                           .replace("（含解析）", "") \
                                           .replace(".docx", "") + "（含答案）.docx"
                elif word == "试卷答案":
                    filter_docx_file = doc_file.replace("（含答案）", "") \
                                           .replace("（含解析）", "") \
                                           .replace(".docx", "") + "（含答案）.docx"
                elif word == "【答案】":
                    filter_docx_file = doc_file.replace("（含答案）", "") \
                                           .replace("（含解析）", "") \
                                           .replace(".docx", "") + "（含答案）.docx"
                if "解析" in para.text and "参考答案" in para.text and "答案要点" in para.text and "试卷答案" in para.text and "【答案】" in para.text:
                    filter_docx_file = doc_file.replace("（含答案）", "") \
                                           .replace("（含解析）", "") \
                                           .replace(".docx", "") + "（含解析）.docx"
                if not os.path.exists(filter_docx_file) and (doc_file != filter_docx_file):
                    os.rename(doc_file, filter_docx_file)
                    return True
                stop = True
                break
        if not stop:
            if file.find("（含解析）") != -1 or file.find("（含答案）") != -1:
                dst_docx_file = doc_file.replace("（含解析）", "").replace("（含答案）", "")
                if not os.path.exists(dst_docx_file):
                    os.rename(doc_file, dst_docx_file)
                    return True
    doc.save(dst_docx_file)


def check_only_image(doc_file):
    doc = Document(doc_file)
    i = 0
    for para in doc.paragraphs:
        if i == 0 and para.text == "":
            doc.save(doc_file)
            return True
        i += 1
    doc.save(doc_file)
    return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            print('转换成功')
            doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


if __name__ == '__main__':
    root_dir = "G:\www.ppt818.com"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)

        # 查找含有需要过滤文字的文件
        # if os.path.splitext(file)[1] == ".docx":
        #     print_filter_word(file_path)

        # 删除页眉页脚
        # if os.path.splitext(file)[1] == ".docx":
        #     remove_header_footer(file_path)

        # 改变文档字体
        # if os.path.splitext(file)[1] == ".docx":
        #     change_word_font(file_path)

        # 移除需要过滤文字
        # if os.path.splitext(file)[1] == ".docx":
        #     docx_remove_content(file_path)

        # 改变文件名称
        # if os.path.splitext(file)[1] == ".docx":
        #     rename_docx_name(file_path)

        # 删除只包含图片
        # if os.path.splitext(file)[1] == ".docx":
        #     if check_only_image(file_path):
        #         os.remove(file_path)

        # 将文件标题“（含答案）”文字移到文件名最后位置
        # year_list = ['2014年', '2015年', '2016年', '2017年', '2018年', '2019年', '2020年', '2021年', '2022年']
        # for year in year_list:
        #     if file.find(year) != -1:
        #         print(file_path)
        #         dst_file = year + file.replace(year, "")
        #         dst_docx_file = root_dir + "\\" + dst_file
        #         if not os.path.exists(dst_docx_file):
        #             os.rename(file_path, dst_docx_file)

        # 将文件标题“（含答案）”文字移到文件名最后位置
        # if file.find("（含解析）") != -1:
        #     print(file_path)
        #     dst_docx_file = file_path.replace(".docx", "") + "（含解析）.docx"
        #     if not os.path.exists(dst_docx_file):
        #         os.rename(file_path, dst_docx_file)
        #     else:
        #         os.remove(file_path)

        # 将文件标题“（含答案）”文字移到文件名最后位置
        # year_list = ['2012年', '2013年', '2014年', '2015年', '2016年', '2017年', '2018年', '2019年', '2020年', '2021年', '2022年']
        # for year in year_list:
        #     if file.find(year) != -1 and (file.find("（含解析）") == -1 and file.find("（含答案）") == -1):
        #         print(file_path)
        #         if os.path.splitext(file)[1] == ".doc":
        #             dst_docx_file = root_dir + "\\" + os.path.splitext(file)[0] + "（含解析）" + ".doc"
        #         elif os.path.splitext(file)[1] == ".doc":
        #             dst_docx_file = root_dir + "\\" + os.path.splitext(file)[0] + "（含解析）" + ".docx"
        #         if not os.path.exists(dst_docx_file):
        #             os.rename(file_path, dst_docx_file)
        #         else:
        #             os.remove(file_path)

        # if os.path.splitext(file)[1] == ".docx":
        #     if file.find("（含答案）") == -1 and file.find("（含解析）") == -1:
        #         print(file_path)
        #         os.remove(file_path)

        dst_file_path = root_dir + "\\" + os.path.splitext(file)[0] + "PPT模板" + ".pptx"
        if not os.path.exists(dst_file_path):
            os.rename(file_path, dst_file_path)

        # 修改文件名称
        # change_name = [
        #     ["PPT模板", ""],
        #     ["PPT模版", ""],
        #     ["PPT", ""],
        #     ["模板", ""],
        #     ["下载", ""],
        # ]
        # for rename in change_name:
        #     if file.find(rename[0]) != -1:
        #         print(file_path)
        #         dst_docx_file = file_path.replace(rename[0], rename[1])
        #         if not os.path.exists(dst_docx_file):
        #             os.rename(file_path, dst_docx_file)
        #         else:
        #             os.remove(file_path)

        # 删除文件
        # delete_name = ["扫描", "原卷版", "图片", "无答案", "自画图", "部分", "仅供参考", "暂缺物理", "(2)", " 2.docx", "听力", "2（"]
        # for dname in delete_name:
        #     if file.find(dname) != -1:
        #         os.remove(file_path)

        # if os.path.splitext(file)[1] in [".doc", ".docx"]:
        #     if os.path.splitext(file)[1] == ".doc":
        #         docx_file = root_dir + "\\" + os.path.splitext(file)[0] + ".docx"
        #         if not os.path.exists(docx_file) and doc2docx(file_path, docx_file):
        #             os.remove(file_path)
        #         else:
        #             os.remove(file_path)
        #         file_path = docx_file
        #
        #     if ".docx" in file_path:
        #         # 删除只包含图片
        #         if check_only_image(file_path):
        #             os.remove(file_path)
        #             continue
        #
        #     if ".docx" in file_path:
        #         # 删除页眉页脚
        #         remove_header_footer(file_path)
