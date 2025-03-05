import re
import time
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.shared import RGBColor  # 设置字体的颜色
from docx.oxml.ns import qn
from win32com import client as wc
from pptx import Presentation
import os
import shutil
import zipfile

zipfile.UNRAR_TOOL = "D:\\Program Files (x86)\\WinRAR\\UnRAR.exe"


def decompress_zip(zip_file_name, dir_name):
    """
    .zip 文件解压
    :param zip_file_name: zip 文件路径
    :param dir_name: 文件解压目录
    :return:
    """
    # 创建 zip 对象
    zip_obj = zipfile.ZipFile(zip_file_name)
    # 目录切换
    if not os.path.exists(dir_name):
        os.mkdir(dir_name)
    os.chdir(dir_name)
    # Extract all files into current directory.
    zip_obj.extractall()
    # zip_obj.extractall(dir_name)
    # 关闭
    zip_obj.close()

def change_word_font(doc_file):
    try:
        # 打开doc文件
        doc = Document(doc_file)
        doc.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 设置中文字体使用字体2->宋体
        doc.save(doc_file)
        return True
    except Exception as e:
        print(e)
        return False


def remove_header_footer(doc):
    # doc：需要去页眉页脚的docx 文件
    try:
        document = Document(doc)
        for section in document.sections:
            section.different_first_page_header_footer = False
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True
        document.save(doc)
        return True
    except Exception as e:
        print(e)
        return False


def doc2docx(in_file, out_file):
    try:
        word = wc.Dispatch("Word.Application")
        try:
            print(in_file)
            print(out_file)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, 12, False, "", True, "", False, False, False, False)
            doc.Close()
            word.Quit()
            return True
        except Exception as e:
            print(e)
    except Exception as e:
        print(e)
    return False


def check_only_image(doc_file):
    try:
        doc = Document(doc_file)
        if len(doc.paragraphs) < 5:
            return True
        else:
            i = 0
            for para in doc.paragraphs:
                if i == 4 and para.text == "":
                    doc.save(doc_file)
                    return True
                i += 1
        doc.save(doc_file)
    except Exception as e:
        print(e)
        return True
    return False


if __name__ == '__main__':
    # 第一步：解压压缩包
    # zip_root_dir = "E:\\workspace\\www.ychedu.com\\2025-03-05\\www.ychedu.com\\中考试题"
    # zip_dirs = sorted(os.listdir(zip_root_dir))
    # zip_files = sorted(os.listdir(zip_root_dir))
    # for zip_file in zip_files:
    #
    #     zip_file_path = zip_root_dir + "\\" + zip_file
    #     dst_file_path = "E:\\workspace\\www.ychedu.com\\2025-03-05\\www.uncompress_ychedu.com\\中考试题"
    #     if not os.path.exists(dst_file_path):
    #         os.makedirs(dst_file_path)
    #     print("==========" + "开始解压" + zip_file_path + "==========")
    #     try:
    #         decompress_zip(zip_file_path, dst_file_path+"\\"+zip_file.replace(".zip", "").replace("下载", "").replace("免费", ""))
    #     except Exception as e:
    #         print(e)
    #         continue
    #     print("==========" + "解压完成" + "==========")
    # exit()

    # 第二步：将文件夹中文件移出，并更改文件名称
    # root_dir = "E:\\workspace\\www.ychedu.com\\2025-03-05\\www.uncompress_ychedu.com\\中考试题"
    # files = sorted(os.listdir(root_dir))
    # for file in files:
    #     file_path = root_dir + "\\" + file
    #     print(file_path)
    #
    #     # 查看一下是否是文件夹，如果是文件夹，则将文件移出
    #     if os.path.isdir(file_path):
    #         child_files = sorted(os.listdir(file_path))
    #         for child_file in child_files:
    #             src_child_file_path = file_path + "\\" + child_file
    #             dst_child_file_path = root_dir + "\\" + child_file
    #             try:
    #                 os.rename(src_child_file_path, dst_child_file_path)
    #             except WindowsError:
    #                 os.remove(dst_child_file_path)
    #                 os.rename(src_child_file_path, dst_child_file_path)
    #
    #     # 文件后缀不是doc或docx，则删除
    #     # if os.path.splitext(file)[1] not in [".doc", ".docx"]:
    #     #     os.remove(file_path)
    # exit()

    # 第三步：将doc文档转化为docx
    root_dir = "E:\\workspace\\www.ychedu.com\\2025-03-05\\www.uncompress_ychedu.com\\中考试题"
    files = sorted(os.listdir(root_dir))
    for file in files:
        file_path = root_dir + "\\" + file
        print(file_path)
        docx_dir = "E:\\workspace\\www.ychedu.com\\2025-03-05\\www.docx_ychedu.com\\中考试题"
        if not os.path.exists(docx_dir):
            os.makedirs(docx_dir)

        sub_file = file
        docx_file = docx_dir + "\\" + sub_file.lower().replace(os.path.splitext(sub_file)[1], ".docx")
        docx_file = docx_file.strip()
        docx_file = docx_file.replace("（", "(").replace("）", ")")
        docx_file = docx_file.replace("word版", "")
        docx_file = docx_file.replace("()", "")
        docx_file = docx_file.replace("【全免费】", "")
        docx_file = docx_file.replace("【中考真题】", "")
        docx_file = docx_file.replace(",", "")
        docx_file = docx_file.replace("，", "")
        print(docx_file)

        if not os.path.exists(docx_file):
            # 获取文件后缀
            file_ext = os.path.splitext(file_path)[-1]
            if file_ext == ".docx":
                # 已经是docx文件了，直接复制过去
                shutil.copy(file_path, docx_file)
            else:
                with open(docx_file, 'w') as f:
                    pass
                print("==========开始转化为docx==============")
                if not doc2docx(file_path, docx_file):
                    # 删除原文件
                    os.remove(docx_file)
                    continue
                print("==========转化完成==============")

        # if os.path.exists(docx_file):
        #     # 删除只包含图片
        #     if check_only_image(docx_file):
        #         # 删除图片文件
        #         os.remove(docx_file)
        #         continue
        #
        #     # 删除页眉页脚
        #     if not remove_header_footer(docx_file):
        #         # 删除原文件
        #         os.remove(docx_file)
        #         continue
        #
        #     # 改变文档字体
        #     if not change_word_font(docx_file):
        #         # 删除原文件
        #         os.remove(docx_file)
        #         continue

        # docx文件已存在，跳过继续
        # if os.path.exists(docx_file):
        #     # continue
        #     finish_dir = "E:\\workspace\\www.ychedu.com\\2025-03-05\\www.docx_ychedu.com\\中考试题"
        #     if not os.path.exists(finish_dir):
        #         os.makedirs(finish_dir)
        #     # 将docx文件转化为pdf
        #     finish_file = docx_file.replace("docx_", "finish_").replace(".docx", ".pdf")
        #     if not os.path.exists(finish_file):
        #         # 将docx转化为pdf
        #         with open(finish_file, "w") as f:
        #             # 将 Word 文档转换为 PDF
        #             try:
        #                 print("==========开始转化为pdf==============")
        #                 convert(docx_file, finish_file)
        #                 print("转换成功！")
        #             except Exception as e:
        #                 print("转换失败：", str(e))
